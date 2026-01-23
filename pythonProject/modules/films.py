# modules/films.py
from fastapi import APIRouter, Request, HTTPException, Depends
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.templating import Jinja2Templates
import database
from contextlib import contextmanager
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os

from modules.auth import get_current_user, check_permission

router = APIRouter()

current_dir = os.path.dirname(os.path.abspath(__file__))
templates_dir = os.path.join(current_dir, "..", "templates")
templates = Jinja2Templates(directory=templates_dir)


@contextmanager
def get_db():
    conn = database.get_db_connection()
    try:
        yield conn
    finally:
        if conn:
            conn.close()


# В функции get_films заменить запрос:
@router.get("/api/films")
async def get_films(current_user: dict = Depends(get_current_user)):
    if not check_permission(current_user, "films:read"):
        raise HTTPException(status_code=403, detail="Доступ запрещен")

    try:
        with get_db() as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    SELECT id_movie, movie_title, movie_description, 
                           duration_minutes, genre, age_rating, 
                           release_year, director, created_at
                    FROM movie 
                    ORDER BY id_movie
                """)
                films_data = cur.fetchall()

                films = []
                for film in films_data:
                    films.append({
                        "id": film["id_movie"],
                        "title": film["movie_title"],
                        "description": film["movie_description"],
                        "duration": film["duration_minutes"],
                        "genre": film["genre"],
                        "age_rating": film["age_rating"],
                        "release_year": film["release_year"],
                        "director": film["director"],
                        "created_at": film["created_at"].strftime("%Y-%m-%d %H:%M") if film["created_at"] else None
                    })

        return films
    except Exception as e:
        return {"error": str(e)}


@router.get("/api/films/{film_id}")
async def get_film(film_id: int, current_user: dict = Depends(get_current_user)):
    if not check_permission(current_user, "films:read"):
        raise HTTPException(status_code=403, detail="Доступ запрещен")

    conn = database.get_db_connection()
    if not conn:
        raise HTTPException(status_code=500, detail="Database connection failed")

    cur = conn.cursor()
    cur.execute("SELECT * FROM movie WHERE id_movie = %s", (film_id,))
    movie = cur.fetchone()
    cur.close()
    conn.close()

    if movie:
        return {
            "id": movie["id_movie"],
            "title": movie["movie_title"],
            "description": movie["movie_description"]
        }
    raise HTTPException(status_code=404, detail="Фильм не найден")


@router.post("/api/films")
async def create_film(film: dict, current_user: dict = Depends(get_current_user)):
    if not check_permission(current_user, "films:create"):
        raise HTTPException(status_code=403, detail="Доступ запрещен")

    conn = database.get_db_connection()
    if not conn:
        raise HTTPException(status_code=500, detail="Database connection failed")

    cur = conn.cursor()
    cur.execute(
        "INSERT INTO movie (movie_title, movie_description) VALUES (%s, %s) RETURNING id_movie",
        (film.get('title'), film.get('description'))
    )
    new_id = cur.fetchone()["id_movie"]
    conn.commit()
    cur.close()
    conn.close()
    return {"message": "Фильм успешно добавлен", "id": new_id}


@router.put("/api/films/{film_id}")
async def update_film(film_id: int, film: dict, current_user: dict = Depends(get_current_user)):
    if not check_permission(current_user, "films:update"):
        raise HTTPException(status_code=403, detail="Доступ запрещен")

    conn = database.get_db_connection()
    if not conn:
        raise HTTPException(status_code=500, detail="Database connection failed")

    cur = conn.cursor()
    cur.execute(
        "UPDATE movie SET movie_title = %s, movie_description = %s WHERE id_movie = %s",
        (film.get('title'), film.get('description'), film_id)
    )
    conn.commit()
    cur.close()
    conn.close()
    return {"message": "Фильм успешно обновлен"}


@router.delete("/api/films/{film_id}")
async def delete_film(film_id: int, current_user: dict = Depends(get_current_user)):
    if not check_permission(current_user, "films:delete"):
        raise HTTPException(status_code=403, detail="Доступ запрещен")

    conn = database.get_db_connection()
    if not conn:
        raise HTTPException(status_code=500, detail="Database connection failed")

    cur = conn.cursor()
    cur.execute("DELETE FROM tickets WHERE id_movie = %s", (film_id,))
    cur.execute("DELETE FROM session WHERE id_movie = %s", (film_id,))
    cur.execute("DELETE FROM movie WHERE id_movie = %s", (film_id,))
    conn.commit()
    cur.close()
    conn.close()
    return {"message": "Фильм успешно удален"}


# ========== ОТЧЕТЫ ==========
@router.get("/films-word-report")
async def films_word_report(current_user: dict = Depends(get_current_user)):
    if not check_permission(current_user, "reports:generate"):
        raise HTTPException(status_code=403, detail="Доступ запрещен")

    try:
        with get_db() as conn:
            with conn.cursor() as cur:
                cur.execute("SELECT * FROM movie ORDER BY id_movie")
                films_data = cur.fetchall()

                films = []
                for film in films_data:
                    films.append({
                        "film_id": film["id_movie"],
                        "title": film["movie_title"],
                        "description": film["movie_description"]
                    })

        # Создаем Word документ
        doc = Document()

        # Заголовок документа
        title = doc.add_heading('КАТАЛОГ ФИЛЬМОВ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Подзаголовок
        doc.add_heading('Отчет о фильмотеке', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        if films:
            # Заголовок таблицы
            doc.add_heading('Список фильмов', level=2)

            # Создаем таблицу
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'

            # Заголовки таблицы
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'ID'
            hdr_cells[1].text = 'Название фильма'
            hdr_cells[2].text = 'Описание'

            # Делаем заголовки жирными
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # Добавляем данные фильмов
            for film in films:
                row_cells = table.add_row().cells
                row_cells[0].text = str(film['film_id'])
                row_cells[1].text = film['title']
                row_cells[2].text = film['description']

            doc.add_paragraph()

            # Итоговая информация
            total_para = doc.add_paragraph()
            total_para.add_run('Всего фильмов в каталоге: ').bold = True
            total_para.add_run(str(len(films)))
        else:
            no_data_para = doc.add_paragraph('В каталоге нет фильмов')
            no_data_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            no_data_para.style.font.italic = True

        # Пустые строки перед подписью
        doc.add_paragraph()
        doc.add_paragraph()

        # Подпись директора
        sign_para = doc.add_paragraph()
        sign_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sign_para.add_run('Директор кинотеатра: ___________________\n')
        sign_para.add_run('Дата: ___________________')

        # Сохраняем документ в память
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        # Возвращаем файл как ответ
        filename = f"films_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    except Exception as e:
        return {"error": f"Ошибка генерации Word документа: {str(e)}"}